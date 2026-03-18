from pathlib import Path
from typing import Callable

from pyconverter.core.base_converter import BaseConverter, ConversionOption, ConversionTask
from pyconverter.core.exceptions import ConversionError


class DocumentConverter(BaseConverter):

    def name(self) -> str:
        return "Document Converter"

    def supported_input_formats(self) -> list[str]:
        return ["pdf", "docx", "txt", "html", "md"]

    def supported_output_formats(self) -> list[str]:
        return ["pdf", "docx", "txt", "html"]

    def get_options(self, input_format: str, output_format: str) -> list[ConversionOption]:
        opts = []
        if output_format == "pdf":
            opts.append(ConversionOption(
                "font_size", "Font Size", "int", 12, min_val=6, max_val=72
            ))
        return opts

    def convert(
        self,
        task: ConversionTask,
        progress_callback: Callable[[float], None] | None = None,
    ) -> Path:
        if progress_callback:
            progress_callback(0.0)

        in_fmt = task.input_path.suffix.lower().lstrip(".")
        out_fmt = task.output_format.lower()

        task.output_path.parent.mkdir(parents=True, exist_ok=True)

        if in_fmt == "txt" and out_fmt == "html":
            self._txt_to_html(task)
        elif in_fmt == "txt" and out_fmt == "pdf":
            self._txt_to_pdf(task)
        elif in_fmt == "txt" and out_fmt == "docx":
            self._txt_to_docx(task)
        elif in_fmt == "md" and out_fmt == "html":
            self._md_to_html(task)
        elif in_fmt == "md" and out_fmt == "pdf":
            self._md_to_pdf(task)
        elif in_fmt == "md" and out_fmt == "txt":
            self._md_to_txt(task)
        elif in_fmt == "md" and out_fmt == "docx":
            # md -> html -> then extract text for docx
            self._md_to_docx(task)
        elif in_fmt == "html" and out_fmt == "txt":
            self._html_to_txt(task)
        elif in_fmt == "html" and out_fmt == "pdf":
            self._html_to_pdf(task)
        elif in_fmt == "html" and out_fmt == "docx":
            self._html_to_docx(task)
        elif in_fmt == "docx" and out_fmt == "txt":
            self._docx_to_txt(task)
        elif in_fmt == "docx" and out_fmt == "html":
            self._docx_to_html(task)
        elif in_fmt == "docx" and out_fmt == "pdf":
            self._docx_to_pdf(task)
        elif in_fmt == "pdf" and out_fmt == "txt":
            self._pdf_to_txt(task)
        elif in_fmt == "pdf" and out_fmt == "html":
            self._pdf_to_html(task)
        elif in_fmt == "pdf" and out_fmt == "docx":
            self._pdf_to_docx(task)
        else:
            raise ConversionError(f"Unsupported conversion: {in_fmt} -> {out_fmt}")

        if progress_callback:
            progress_callback(1.0)

        return task.output_path

    # --- TXT conversions ---

    def _txt_to_html(self, task: ConversionTask):
        text = task.input_path.read_text(encoding="utf-8")
        html = f"<html><body><pre>{text}</pre></body></html>"
        task.output_path.write_text(html, encoding="utf-8")

    def _txt_to_pdf(self, task: ConversionTask):
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        text = task.input_path.read_text(encoding="utf-8")
        font_size = task.options.get("font_size", 12)

        c = canvas.Canvas(str(task.output_path), pagesize=A4)
        width, height = A4
        y = height - 40
        for line in text.splitlines():
            if y < 40:
                c.showPage()
                y = height - 40
            c.setFont("Helvetica", font_size)
            c.drawString(40, y, line)
            y -= font_size + 4
        c.save()

    def _txt_to_docx(self, task: ConversionTask):
        from docx import Document

        text = task.input_path.read_text(encoding="utf-8")
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        doc.save(str(task.output_path))

    # --- Markdown conversions ---

    def _md_to_html(self, task: ConversionTask):
        import markdown

        md_text = task.input_path.read_text(encoding="utf-8")
        html_body = markdown.markdown(md_text, extensions=["tables", "fenced_code"])
        html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"></head>
<body>{html_body}</body></html>"""
        task.output_path.write_text(html, encoding="utf-8")

    def _md_to_pdf(self, task: ConversionTask):
        import markdown
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        md_text = task.input_path.read_text(encoding="utf-8")
        # Convert to plain text for PDF (strip markdown)
        import re
        plain = re.sub(r"[#*_`~\[\]()]", "", md_text)

        font_size = task.options.get("font_size", 12)
        c = canvas.Canvas(str(task.output_path), pagesize=A4)
        width, height = A4
        y = height - 40
        for line in plain.splitlines():
            if y < 40:
                c.showPage()
                y = height - 40
            c.setFont("Helvetica", font_size)
            c.drawString(40, y, line)
            y -= font_size + 4
        c.save()

    def _md_to_txt(self, task: ConversionTask):
        import re
        md_text = task.input_path.read_text(encoding="utf-8")
        plain = re.sub(r"[#*_`~\[\]()]", "", md_text)
        task.output_path.write_text(plain, encoding="utf-8")

    def _md_to_docx(self, task: ConversionTask):
        import re
        from docx import Document

        md_text = task.input_path.read_text(encoding="utf-8")
        plain = re.sub(r"[#*_`~\[\]()]", "", md_text)
        doc = Document()
        for line in plain.splitlines():
            doc.add_paragraph(line)
        doc.save(str(task.output_path))

    # --- HTML conversions ---

    def _html_to_txt(self, task: ConversionTask):
        import re
        html = task.input_path.read_text(encoding="utf-8")
        text = re.sub(r"<[^>]+>", "", html)
        task.output_path.write_text(text.strip(), encoding="utf-8")

    def _html_to_pdf(self, task: ConversionTask):
        import re
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        html = task.input_path.read_text(encoding="utf-8")
        text = re.sub(r"<[^>]+>", "", html)
        font_size = task.options.get("font_size", 12)

        c = canvas.Canvas(str(task.output_path), pagesize=A4)
        width, height = A4
        y = height - 40
        for line in text.strip().splitlines():
            if y < 40:
                c.showPage()
                y = height - 40
            c.setFont("Helvetica", font_size)
            c.drawString(40, y, line)
            y -= font_size + 4
        c.save()

    def _html_to_docx(self, task: ConversionTask):
        import re
        from docx import Document

        html = task.input_path.read_text(encoding="utf-8")
        text = re.sub(r"<[^>]+>", "", html)
        doc = Document()
        for line in text.strip().splitlines():
            doc.add_paragraph(line)
        doc.save(str(task.output_path))

    # --- DOCX conversions ---

    def _docx_to_txt(self, task: ConversionTask):
        from docx import Document

        doc = Document(str(task.input_path))
        text = "\n".join(p.text for p in doc.paragraphs)
        task.output_path.write_text(text, encoding="utf-8")

    def _docx_to_html(self, task: ConversionTask):
        from docx import Document

        doc = Document(str(task.input_path))
        paragraphs = [f"<p>{p.text}</p>" for p in doc.paragraphs]
        html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"></head>
<body>{"".join(paragraphs)}</body></html>"""
        task.output_path.write_text(html, encoding="utf-8")

    def _docx_to_pdf(self, task: ConversionTask):
        from docx import Document
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        doc = Document(str(task.input_path))
        font_size = task.options.get("font_size", 12)

        c = canvas.Canvas(str(task.output_path), pagesize=A4)
        width, height = A4
        y = height - 40
        for para in doc.paragraphs:
            if y < 40:
                c.showPage()
                y = height - 40
            c.setFont("Helvetica", font_size)
            c.drawString(40, y, para.text)
            y -= font_size + 4
        c.save()

    # --- PDF conversions ---

    def _pdf_to_txt(self, task: ConversionTask):
        import fitz

        doc = fitz.open(str(task.input_path))
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        task.output_path.write_text("\n".join(text_parts), encoding="utf-8")

    def _pdf_to_html(self, task: ConversionTask):
        import fitz

        doc = fitz.open(str(task.input_path))
        html_parts = ['<!DOCTYPE html><html><head><meta charset="utf-8"></head><body>']
        for page in doc:
            text = page.get_text()
            paragraphs = [f"<p>{line}</p>" for line in text.splitlines() if line.strip()]
            html_parts.extend(paragraphs)
            html_parts.append("<hr>")
        doc.close()
        html_parts.append("</body></html>")
        task.output_path.write_text("\n".join(html_parts), encoding="utf-8")

    def _pdf_to_docx(self, task: ConversionTask):
        import fitz
        from docx import Document

        pdf = fitz.open(str(task.input_path))
        doc = Document()
        for page in pdf:
            text = page.get_text()
            for line in text.splitlines():
                if line.strip():
                    doc.add_paragraph(line)
        pdf.close()
        doc.save(str(task.output_path))
